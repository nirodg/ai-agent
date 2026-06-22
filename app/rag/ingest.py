"""Document ingestion: load files / external sources, chunk, embed into a project."""

import csv
import io
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.db import add_rag_document
from app.rag.store import add_documents

_SPLITTER = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)


def _split(text: str, source: str, extra_meta: dict | None = None) -> List[Document]:
    meta = {"source": source}
    if extra_meta:
        meta.update(extra_meta)
    chunks = _SPLITTER.split_text(text)
    return [Document(page_content=c, metadata=dict(meta)) for c in chunks if c.strip()]


def _read_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)


def _read_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    if not rows:
        return ""
    header = rows[0]
    lines = []
    for row in rows[1:]:
        pairs = [f"{h}: {v}" for h, v in zip(header, row)]
        lines.append("; ".join(pairs))
    return "\n".join(lines)


def ingest_file(project_id: int, filename: str, data: bytes) -> int:
    """Ingest an uploaded file into a project's knowledge base.

    Supports PDF, DOCX, CSV, TXT/MD. Returns number of chunks stored.
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf(data)
        source_type = "pdf"
    elif suffix in (".docx", ".doc"):
        text = _read_docx(data)
        source_type = "docx"
    elif suffix == ".csv":
        text = _read_csv(data)
        source_type = "csv"
    elif suffix in (".txt", ".md", ".markdown"):
        text = data.decode("utf-8", errors="replace")
        source_type = "text"
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    docs = _split(text, source=filename, extra_meta={"type": source_type})
    count = add_documents(project_id, docs)
    add_rag_document(project_id, filename, source_type, count)
    return count


def ingest_text(project_id: int, source: str, text: str, source_type: str = "text") -> int:
    """Ingest raw text (e.g. from a connector) into a project."""
    docs = _split(text, source=source, extra_meta={"type": source_type})
    count = add_documents(project_id, docs)
    add_rag_document(project_id, source, source_type, count)
    return count
